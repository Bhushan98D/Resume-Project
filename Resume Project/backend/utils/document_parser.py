"""
Document parsing utilities for PDF and DOCX files.
"""

import os
import json
import re
from typing import Dict, List, Optional, Tuple
import fitz  # PyMuPDF
import pdfplumber
import docx
import docx2txt
from pathlib import Path


class DocumentParser:
    """Parser for extracting text from PDF and DOCX documents."""
    
    def __init__(self):
        self.supported_formats = ['.pdf', '.docx']
    
    def parse_document(self, file_path: str) -> Dict:
        """
        Parse document and extract text with metadata.
        
        Args:
            file_path: Path to the document file
            
        Returns:
            Dictionary containing extracted text and metadata
        """
        file_path = Path(file_path)
        
        if not file_path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")
        
        file_extension = file_path.suffix.lower()
        
        if file_extension not in self.supported_formats:
            raise ValueError(f"Unsupported file format: {file_extension}")
        
        if file_extension == '.pdf':
            return self._parse_pdf(file_path)
        elif file_extension == '.docx':
            return self._parse_docx(file_path)
    
    def _parse_pdf(self, file_path: Path) -> Dict:
        """Parse PDF file using both PyMuPDF and pdfplumber for better extraction."""
        text_content = []
        metadata = {}
        
        try:
            # Try PyMuPDF first
            doc = fitz.open(file_path)
            metadata = doc.metadata
            text_content = []
            
            for page_num in range(len(doc)):
                page = doc.load_page(page_num)
                text = page.get_text()
                text_content.append(text)
            
            doc.close()
            
            # Fallback to pdfplumber if PyMuPDF fails
            if not text_content or len(' '.join(text_content).strip()) < 100:
                with pdfplumber.open(file_path) as pdf:
                    text_content = []
                    for page in pdf.pages:
                        text = page.extract_text()
                        if text:
                            text_content.append(text)
        
        except Exception as e:
            print(f"Error parsing PDF with PyMuPDF: {e}")
            # Fallback to pdfplumber
            try:
                with pdfplumber.open(file_path) as pdf:
                    text_content = []
                    for page in pdf.pages:
                        text = page.extract_text()
                        if text:
                            text_content.append(text)
            except Exception as e2:
                raise Exception(f"Failed to parse PDF with both methods: {e2}")
        
        full_text = '\n'.join(text_content)
        cleaned_text = self._clean_text(full_text)
        
        return {
            'text': cleaned_text,
            'raw_text': full_text,
            'metadata': metadata,
            'file_type': 'PDF',
            'page_count': len(text_content)
        }
    
    def _parse_docx(self, file_path: Path) -> Dict:
        """Parse DOCX file using python-docx."""
        try:
            # Try python-docx first
            doc = docx.Document(file_path)
            text_content = []
            
            for paragraph in doc.paragraphs:
                text_content.append(paragraph.text)
            
            # Also extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text_content.append(cell.text)
            
            full_text = '\n'.join(text_content)
            
        except Exception as e:
            print(f"Error parsing DOCX with python-docx: {e}")
            # Fallback to docx2txt
            try:
                full_text = docx2txt.process(file_path)
            except Exception as e2:
                raise Exception(f"Failed to parse DOCX with both methods: {e2}")
        
        cleaned_text = self._clean_text(full_text)
        
        return {
            'text': cleaned_text,
            'raw_text': full_text,
            'metadata': {},
            'file_type': 'DOCX',
            'page_count': 1  # DOCX doesn't have clear page count
        }
    
    def _clean_text(self, text: str) -> str:
        """Clean and normalize extracted text."""
        if not text:
            return ""
        
        # Remove excessive whitespace
        text = re.sub(r'\s+', ' ', text)
        
        # Remove common headers/footers patterns
        lines = text.split('\n')
        cleaned_lines = []
        
        for line in lines:
            line = line.strip()
            if not line:
                continue
            
            # Skip common header/footer patterns
            if self._is_header_footer(line):
                continue
            
            cleaned_lines.append(line)
        
        return '\n'.join(cleaned_lines)
    
    def _is_header_footer(self, line: str) -> bool:
        """Check if line is likely a header or footer."""
        line_lower = line.lower()
        
        # Common header/footer patterns
        patterns = [
            r'^\d+$',  # Page numbers
            r'^page \d+ of \d+$',
            r'^confidential$',
            r'^draft$',
            r'^internal use only$',
            r'^\d{1,2}/\d{1,2}/\d{4}$',  # Dates
            r'^\d{4}-\d{2}-\d{2}$',  # ISO dates
        ]
        
        for pattern in patterns:
            if re.match(pattern, line_lower):
                return True
        
        # Check for very short lines that might be headers/footers
        if len(line) < 3 and line.isdigit():
            return True
        
        return False
    
    def extract_sections(self, text: str) -> Dict[str, str]:
        """Extract different sections from resume text."""
        sections = {
            'contact_info': '',
            'summary': '',
            'experience': '',
            'education': '',
            'skills': '',
            'projects': '',
            'certifications': '',
            'other': ''
        }
        
        # Common section headers
        section_patterns = {
            'contact_info': r'(contact|personal|address|phone|email)',
            'summary': r'(summary|profile|objective|about)',
            'experience': r'(experience|work history|employment|career)',
            'education': r'(education|academic|qualification)',
            'skills': r'(skills|technical skills|competencies)',
            'projects': r'(projects|portfolio)',
            'certifications': r'(certifications|certificates|licenses)'
        }
        
        lines = text.split('\n')
        current_section = 'other'
        
        for line in lines:
            line = line.strip()
            if not line:
                continue
            
            # Check if this line is a section header
            for section, pattern in section_patterns.items():
                if re.search(pattern, line, re.IGNORECASE):
                    current_section = section
                    break
            
            # Add line to current section
            if current_section in sections:
                if sections[current_section]:
                    sections[current_section] += '\n' + line
                else:
                    sections[current_section] = line
        
        return sections
    
    def extract_skills(self, text: str) -> List[str]:
        """Extract skills from resume text."""
        # Common technical skills patterns
        skill_patterns = [
            r'\b(?:python|java|javascript|typescript|react|angular|vue|node\.?js|express|django|flask|fastapi)\b',
            r'\b(?:sql|mysql|postgresql|mongodb|redis|elasticsearch)\b',
            r'\b(?:aws|azure|gcp|docker|kubernetes|terraform)\b',
            r'\b(?:git|github|gitlab|jenkins|ci/cd)\b',
            r'\b(?:machine learning|ml|ai|deep learning|nlp|computer vision)\b',
            r'\b(?:pandas|numpy|scikit-learn|tensorflow|pytorch)\b',
            r'\b(?:html|css|bootstrap|sass|less)\b',
            r'\b(?:agile|scrum|kanban|devops)\b'
        ]
        
        skills = set()
        text_lower = text.lower()
        
        for pattern in skill_patterns:
            matches = re.findall(pattern, text_lower, re.IGNORECASE)
            skills.update(matches)
        
        return list(skills)
    
    def extract_education(self, text: str) -> List[Dict]:
        """Extract education information from resume text."""
        education = []
        
        # Common degree patterns
        degree_patterns = [
            r'\b(?:bachelor|master|phd|doctorate|diploma|certificate)\b.*?(?:in|of)\s+([^,\n]+)',
            r'\b(?:b\.?s\.?|m\.?s\.?|ph\.?d\.?|mba|bca|mca)\b.*?(?:in|of)\s+([^,\n]+)',
        ]
        
        lines = text.split('\n')
        for line in lines:
            line = line.strip()
            for pattern in degree_patterns:
                match = re.search(pattern, line, re.IGNORECASE)
                if match:
                    education.append({
                        'degree': match.group(0),
                        'field': match.group(1).strip() if len(match.groups()) > 0 else '',
                        'line': line
                    })
        
        return education


# Example usage and testing
if __name__ == "__main__":
    parser = DocumentParser()
    
    # Test with a sample file (if available)
    test_file = "sample_resume.pdf"
    if os.path.exists(test_file):
        result = parser.parse_document(test_file)
        print("Extracted text length:", len(result['text']))
        print("File type:", result['file_type'])
        
        sections = parser.extract_sections(result['text'])
        print("\nSections found:")
        for section, content in sections.items():
            if content:
                print(f"{section}: {len(content)} characters")
        
        skills = parser.extract_skills(result['text'])
        print(f"\nSkills found: {skills}")
        
        education = parser.extract_education(result['text'])
        print(f"\nEducation found: {education}")
    else:
        print("No test file found. Please add a sample resume to test.")
